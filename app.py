import streamlit as st
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from docx.shared import RGBColor
import fitz  # PyMuPDF for PDF processing
import io
import pymupdf 

st.set_page_config(page_title="Plagiarism Detection", page_icon="⚖️")

@st.cache_resource
def load_model():
    local_model_path = "all-MiniLM-L6-v2"  # Path to local model or download from Hugging Face
    model = SentenceTransformer(local_model_path)
    return model

model = load_model()

def extract_text_from_pdf(uploaded_file):
    text = ""
    doc = pymupdf .open(stream=uploaded_file.getvalue(), filetype="pdf")  # Corrected method
    for page in doc:
        text += page.get_text("text")
    return text

def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_txt(uploaded_file):
    return uploaded_file.read().decode("utf-8")

def calculate_similarity(text1, text2):
    embedding1 = model.encode(text1, convert_to_tensor=True)
    embedding2 = model.encode(text2, convert_to_tensor=True)
    similarity = cosine_similarity([embedding1.cpu().numpy()], [embedding2.cpu().numpy()])
    return similarity[0][0]

def compare_files(files):
    file_texts = []
    for file in files:
        file_name = file.name
        if file_name.endswith(".pdf"):
            text = extract_text_from_pdf(file)
        elif file_name.endswith(".docx"):
            text = extract_text_from_docx(file)
        elif file_name.endswith(".txt"):
            text = extract_text_from_txt(file)
        else:
            continue
        file_texts.append((text, file))
    return file_texts

def display_plagiarism_results(file_texts):
    st.write("### Plagiarism Detection Results")
    plagiarism_detected = False
    first_detection = False
    
    for idx1, (text1, file1) in enumerate(file_texts):
        for idx2, (text2, file2) in enumerate(file_texts):
            if idx1 < idx2:
                similarity = calculate_similarity(text1, text2)
                
                if similarity > 0.8:
                    plagiarism_detected = True
                    st.error(f"Plagiarism Detected between {file1.name} and {file2.name} - Similarity: {similarity*100:.2f}%")
                    
                    if not first_detection:
                        first_detection = True
                        document = Document()
                        document.add_heading('Plagiarism Detected', 0)
                        
                        p = document.add_paragraph()
                        run = p.add_run(f"Plagiarism between {file1.name} and {file2.name}. Similarity: {similarity*100:.2f}%")
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        p.add_run("\n\n")

                        p.add_run(f"Content from {file1.name}:\n").bold = True
                        p.add_run(text1[:500] + "...\n\n")

                        p.add_run(f"Content from {file2.name}:\n").bold = True
                        p.add_run(text2[:500] + "...\n\n")

                        byte_io = io.BytesIO()
                        document.save(byte_io)
                        byte_io.seek(0)

                        st.download_button(
                            label="Download Plagiarism Report",
                            data=byte_io,
                            file_name="Detected.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

    if plagiarism_detected:
        st.markdown("<p style='color:red; font-size:20px;'>Plagiarism detected.</p>", unsafe_allow_html=True)
    else:
        st.markdown("<p style='color:green; font-size:20px;'>No plagiarism detected.</p>", unsafe_allow_html=True)

st.title("Plagiarism Detection Tool")
st.write("Upload at least two files (PDF, DOCX, or TXT) to check for plagiarism.")

uploaded_files = st.file_uploader("Choose files", accept_multiple_files=True, type=['pdf', 'docx', 'txt'])

if len(uploaded_files) >= 2:
    file_texts = compare_files(uploaded_files)
    display_plagiarism_results(file_texts)
else:
    st.warning("Please upload at least two files for plagiarism detection.")